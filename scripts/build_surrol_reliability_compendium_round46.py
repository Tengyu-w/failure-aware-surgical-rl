# -*- coding: utf-8 -*-
"""Build a Chinese DOCX compendium for the constraint surgical RL project.

This script intentionally summarizes completed experiment evidence. It does not
modify raw logs, checkpoints, datasets, or run directories.
"""

from __future__ import annotations

from pathlib import Path
from typing import Iterable, Sequence

from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
REPORTS = ROOT / "reports"
FIG_DIR = REPORTS / "figures" / "round46_compendium"
OUT = REPORTS / "constraint_surgical_rl_surrol_reliability_compendium_round46.docx"


BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
INK = "0B2545"
MUTED = "5B6470"
LIGHT_BLUE = "E8EEF5"
LIGHT_GRAY = "F2F4F7"
PALE_GREEN = "E7F4EA"
PALE_GOLD = "FFF4D6"
PALE_RED = "FDE7E9"
BORDER = "C8D1DC"
WHITE = "FFFFFF"


def font_path() -> str | None:
    candidates = [
        r"C:\Windows\Fonts\msyh.ttc",
        r"C:\Windows\Fonts\simhei.ttf",
        r"C:\Windows\Fonts\arial.ttf",
    ]
    for p in candidates:
        if Path(p).exists():
            return p
    return None


FONT_PATH = font_path()


def pil_font(size: int, bold: bool = False) -> ImageFont.ImageFont:
    if FONT_PATH:
        return ImageFont.truetype(FONT_PATH, size=size)
    return ImageFont.load_default()


def ensure_dirs() -> None:
    FIG_DIR.mkdir(parents=True, exist_ok=True)


def draw_wrapped(
    draw: ImageDraw.ImageDraw,
    text: str,
    xy: tuple[int, int],
    font: ImageFont.ImageFont,
    fill: str,
    max_width: int,
    line_gap: int = 5,
) -> int:
    x, y = xy
    lines: list[str] = []
    current = ""
    for ch in text:
        probe = current + ch
        if draw.textbbox((0, 0), probe, font=font)[2] <= max_width or not current:
            current = probe
        else:
            lines.append(current)
            current = ch
    if current:
        lines.append(current)
    for line in lines:
        draw.text((x, y), line, font=font, fill=fill)
        y += font.size + line_gap
    return y


def save_stage_overview() -> Path:
    path = FIG_DIR / "stage_status_overview.png"
    w, h = 1500, 760
    im = Image.new("RGB", (w, h), f"#{WHITE}")
    d = ImageDraw.Draw(im)
    title = pil_font(34)
    label = pil_font(23)
    small = pil_font(19)
    d.text((55, 35), "项目阶段总览：从 3D proxy 到 SurRoL 可靠性监督", font=title, fill=f"#{INK}")
    stages = [
        ("旧版基础", "3D proxy\nPPO + tangent shield", "完成"),
        ("Stage 1", "锁定 SurRoL\n视觉风险主线", "完成"),
        ("Stage 2", "真实在线\nadapter 数据", "完成"),
        ("Stage 3", "在线 KNN memory", "完成但负结果"),
        ("Stage 4", "保守风险路由\nbudget 10", "完成"),
        ("Stage 5", "跨任务广度", "有限完成"),
        ("Stage 6", "博士申请叙事", "完成"),
    ]
    x0, y0, box_w, box_h, gap = 55, 145, 190, 190, 18
    colors = [LIGHT_BLUE, LIGHT_BLUE, PALE_GREEN, PALE_GOLD, PALE_GREEN, PALE_GOLD, LIGHT_BLUE]
    for i, (name, desc, status) in enumerate(stages):
        x = x0 + i * (box_w + gap)
        d.rounded_rectangle((x, y0, x + box_w, y0 + box_h), radius=18, fill=f"#{colors[i]}", outline=f"#{BORDER}", width=3)
        d.text((x + 18, y0 + 18), name, font=label, fill=f"#{DARK_BLUE}")
        yy = y0 + 62
        for line in desc.splitlines():
            d.text((x + 18, yy), line, font=small, fill=f"#{INK}")
            yy += 29
        d.rounded_rectangle((x + 18, y0 + 143, x + box_w - 18, y0 + 174), radius=8, fill=f"#{WHITE}", outline=f"#{BORDER}")
        d.text((x + 30, y0 + 147), status, font=small, fill=f"#{MUTED}")
        if i < len(stages) - 1:
            ax = x + box_w + 2
            ay = y0 + 92
            d.line((ax, ay, ax + gap - 4, ay), fill=f"#{BORDER}", width=4)
            d.polygon([(ax + gap - 4, ay), (ax + gap - 16, ay - 8), (ax + gap - 16, ay + 8)], fill=f"#{BORDER}")
    notes = [
        "绿色：可作为当前正向证据；黄色：完成但需作为限制或负结果解释。",
        "核心升级不是替换主控制器，而是在外部增加 failure-aware / risk-aware supervisor。",
    ]
    y = 415
    for note in notes:
        y = draw_wrapped(d, note, (70, y), small, f"#{INK}", 1320, 6) + 6
    im.save(path)
    return path


def save_pipeline() -> Path:
    path = FIG_DIR / "system_pipeline.png"
    w, h = 1500, 840
    im = Image.new("RGB", (w, h), f"#{WHITE}")
    d = ImageDraw.Draw(im)
    title = pil_font(34)
    label = pil_font(23)
    small = pil_font(18)
    d.text((55, 35), "外接式可靠性监督管线", font=title, fill=f"#{INK}")
    boxes = [
        (80, 150, 330, 285, "SurRoL 观测", "proprioception + RGB pooled vision\n208D render_proprio_vision"),
        (430, 150, 680, 285, "主策略", "NeedleReach visual DAgger policy\n输出 5D action"),
        (780, 150, 1030, 285, "风险与记忆", "learned risk head\nrecovery memory / OOD distance"),
        (1130, 150, 1380, 285, "路由", "auto / recovery / review\nabort_candidate"),
    ]
    for x1, y1, x2, y2, head, body in boxes:
        d.rounded_rectangle((x1, y1, x2, y2), radius=18, fill=f"#{LIGHT_BLUE}", outline=f"#{BORDER}", width=3)
        d.text((x1 + 18, y1 + 18), head, font=label, fill=f"#{DARK_BLUE}")
        yy = y1 + 58
        for line in body.splitlines():
            d.text((x1 + 18, yy), line, font=small, fill=f"#{INK}")
            yy += 27
    for ax in [340, 690, 1040]:
        d.line((ax, 218, ax + 80, 218), fill=f"#{BORDER}", width=5)
        d.polygon([(ax + 80, 218), (ax + 64, 207), (ax + 64, 229)], fill=f"#{BORDER}")
    routes = [
        ("auto_execute", "低风险，允许继续自动执行", PALE_GREEN),
        ("auto_recovery", "短程、预算内恢复", LIGHT_BLUE),
        ("human_review", "高风险/OOD/停滞，交给人工复核", PALE_GOLD),
        ("abort_candidate", "不可逆风险候选，不盲目 recovery", PALE_RED),
    ]
    y = 410
    for name, desc, fill in routes:
        d.rounded_rectangle((210, y, 1290, y + 72), radius=14, fill=f"#{fill}", outline=f"#{BORDER}", width=2)
        d.text((235, y + 18), name, font=label, fill=f"#{DARK_BLUE}")
        d.text((520, y + 20), desc, font=small, fill=f"#{INK}")
        y += 92
    d.text((80, 785), "定位：尽量不改变主流手术机器人控制逻辑，而是在外部判断什么时候不能继续自动恢复。", font=small, fill=f"#{MUTED}")
    im.save(path)
    return path


def save_bar_chart(
    path: Path,
    title_text: str,
    rows: Sequence[tuple[str, Sequence[float]]],
    labels: Sequence[str],
    colors: Sequence[str],
    max_value: float,
    value_suffix: str = "",
) -> Path:
    w = 1500
    row_h = 74
    top = 130
    bottom = 80
    h = top + len(rows) * row_h + bottom
    im = Image.new("RGB", (w, h), f"#{WHITE}")
    d = ImageDraw.Draw(im)
    title = pil_font(34)
    font = pil_font(19)
    small = pil_font(17)
    d.text((55, 35), title_text, font=title, fill=f"#{INK}")
    plot_x = 420
    plot_w = 850
    legend_x = 1080
    for i, (lab, col) in enumerate(zip(labels, colors)):
        x = legend_x
        y = 48 + i * 34
        d.rectangle((x, y, x + 24, y + 20), fill=f"#{col}", outline=f"#{BORDER}")
        d.text((x + 34, y - 2), lab, font=small, fill=f"#{INK}")
    for ridx, (name, vals) in enumerate(rows):
        y = top + ridx * row_h
        d.text((55, y + 8), name, font=font, fill=f"#{INK}")
        x = plot_x
        for v, c in zip(vals, colors):
            bw = int(plot_w * (v / max_value))
            d.rounded_rectangle((x, y + 11, x + bw, y + 40), radius=5, fill=f"#{c}")
            d.text((x + bw + 8, y + 12), f"{v:g}{value_suffix}", font=small, fill=f"#{MUTED}")
            x += bw + 58
        d.line((plot_x, y + 58, plot_x + plot_w + 170, y + 58), fill="#E7ECF2", width=1)
    im.save(path)
    return path


def make_figures() -> dict[str, Path]:
    ensure_dirs()
    figures: dict[str, Path] = {}
    figures["stage"] = save_stage_overview()
    figures["pipeline"] = save_pipeline()
    figures["dataset"] = save_bar_chart(
        FIG_DIR / "online_adapter_dataset_counts.png",
        "Round 42：真实在线 adapter-space 数据规模",
        [
            ("mixed p=0.35", [970, 596]),
            ("clean", [943, 546]),
            ("total", [1913, 1142]),
        ],
        ["steps", "high action-gap"],
        ["4C78A8", "F58518"],
        2000,
    )
    figures["memory"] = save_bar_chart(
        FIG_DIR / "recovery_memory_comparison_round43.png",
        "Round 43：Recovery memory 闭环对比，负结果必须保留",
        [
            ("old augmented memory", [6, 2, 12]),
            ("Round 40 bad memory", [4, 9, 7]),
            ("Round 43 online KNN", [2, 0, 18]),
        ],
        ["auto success", "auto fail", "review"],
        ["54A24B", "E45756", "B279A2"],
        20,
    )
    figures["routing"] = save_bar_chart(
        FIG_DIR / "conservative_routing_round44.png",
        "Round 44：保守路由把自动失败压到 0",
        [
            ("baseline old risk + old memory", [6, 2, 12]),
            ("+ recovery budget 10", [6, 0, 14]),
            ("+ learned stagnation", [6, 0, 14]),
            ("new risk + old memory", [4, 0, 16]),
        ],
        ["auto success", "auto fail", "review"],
        ["54A24B", "E45756", "B279A2"],
        20,
    )
    figures["cross"] = save_bar_chart(
        FIG_DIR / "cross_task_probe_round45.png",
        "Round 45：learned visual 直接跨任务迁移未成立",
        [
            ("NeedlePick transfer probe", [0, 5]),
            ("GauzeRetrieve transfer probe", [0, 5]),
            ("NeedlePick phase-aware", [10, 0]),
            ("GauzeRetrieve phase-aware", [10, 0]),
        ],
        ["success/recovered", "review/fail"],
        ["54A24B", "B279A2"],
        10,
    )
    return figures


def set_east_asia(run, font_name: str = "Microsoft YaHei") -> None:
    run.font.name = font_name
    if run._element.rPr is None:
        run._element.get_or_add_rPr()
    run._element.rPr.rFonts.set(qn("w:ascii"), font_name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), font_name)
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)


def set_paragraph_font(paragraph, size: float | None = None, color: str | None = None, bold: bool | None = None) -> None:
    for run in paragraph.runs:
        set_east_asia(run)
        if size is not None:
            run.font.size = Pt(size)
        if color is not None:
            run.font.color.rgb = RGBColor.from_string(color)
        if bold is not None:
            run.bold = bold


def style_doc(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)
    section.different_first_page_header_footer = True

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Microsoft YaHei"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(10.5)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    for name, size, color, before, after in [
        ("Heading 1", 16, BLUE, 18, 10),
        ("Heading 2", 13, BLUE, 14, 7),
        ("Heading 3", 12, DARK_BLUE, 10, 5),
    ]:
        st = styles[name]
        st.font.name = "Microsoft YaHei"
        st._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        st.font.size = Pt(size)
        st.font.color.rgb = RGBColor.from_string(color)
        st.font.bold = True
        st.paragraph_format.space_before = Pt(before)
        st.paragraph_format.space_after = Pt(after)
        st.paragraph_format.keep_with_next = True

    header = section.header.paragraphs[0]
    header.text = "Constraint Surgical RL | Failure-aware reliability supervisor"
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    set_paragraph_font(header, 8.5, MUTED)
    footer = section.footer.paragraphs[0]
    footer.text = "Simulation-only research prototype. Not clinical deployment evidence."
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_font(footer, 8.5, MUTED)


def shade_cell(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in [("top", top), ("start", start), ("bottom", bottom), ("end", end)]:
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_table_width(table, widths: Sequence[float]) -> None:
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    for row in table.rows:
        for idx, width in enumerate(widths):
            row.cells[idx].width = Inches(width)
            set_cell_margins(row.cells[idx])
            row.cells[idx].vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def add_table(doc: Document, headers: Sequence[str], rows: Sequence[Sequence[str]], widths: Sequence[float]) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    set_table_width(table, widths)
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        shade_cell(cell, LIGHT_BLUE)
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            set_paragraph_font(p, 8.8, INK, True)
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            cells[i].text = str(value)
            shade_cell(cells[i], WHITE)
            for p in cells[i].paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT if len(str(value)) > 14 else WD_ALIGN_PARAGRAPH.CENTER
                set_paragraph_font(p, 8.4, INK)
    set_table_width(table, widths)
    doc.add_paragraph()


def add_para(doc: Document, text: str, bold: bool = False, color: str = INK, size: float = 10.5) -> None:
    p = doc.add_paragraph()
    r = p.add_run(text)
    set_east_asia(r)
    r.font.size = Pt(size)
    r.font.color.rgb = RGBColor.from_string(color)
    r.bold = bold


def add_bullets(doc: Document, items: Iterable[str]) -> None:
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        r = p.add_run(item)
        set_east_asia(r)
        r.font.size = Pt(10.2)
        r.font.color.rgb = RGBColor.from_string(INK)


def add_numbered(doc: Document, items: Iterable[str]) -> None:
    for item in items:
        p = doc.add_paragraph(style="List Number")
        r = p.add_run(item)
        set_east_asia(r)
        r.font.size = Pt(10.2)
        r.font.color.rgb = RGBColor.from_string(INK)


def add_callout(doc: Document, title: str, body: str, fill: str = LIGHT_GRAY) -> None:
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    set_table_width(table, [6.5])
    cell = table.cell(0, 0)
    shade_cell(cell, fill)
    p = cell.paragraphs[0]
    r = p.add_run(title + "：")
    set_east_asia(r)
    r.font.bold = True
    r.font.size = Pt(10.5)
    r.font.color.rgb = RGBColor.from_string(DARK_BLUE)
    r2 = p.add_run(body)
    set_east_asia(r2)
    r2.font.size = Pt(10.2)
    r2.font.color.rgb = RGBColor.from_string(INK)
    doc.add_paragraph()


def add_figure(doc: Document, path: Path, caption: str, width: float = 6.25) -> None:
    if not path.exists():
        add_callout(doc, "缺失图", f"{path}", PALE_RED)
        return
    doc.add_picture(str(path), width=Inches(width))
    last = doc.paragraphs[-1]
    last.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = cap.add_run(caption)
    set_east_asia(r)
    r.font.size = Pt(8.8)
    r.font.color.rgb = RGBColor.from_string(MUTED)


def add_cover(doc: Document) -> None:
    for _ in range(4):
        doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Constraint Surgical RL")
    set_east_asia(r)
    r.font.size = Pt(28)
    r.font.bold = True
    r.font.color.rgb = RGBColor.from_string(INK)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("SurRoL 可靠性监督实验全册")
    set_east_asia(r)
    r.font.size = Pt(22)
    r.font.bold = True
    r.font.color.rgb = RGBColor.from_string(BLUE)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("视觉不确定性、失败恢复、风险分流与博士申请 Demo 证据整理")
    set_east_asia(r)
    r.font.size = Pt(12)
    r.font.color.rgb = RGBColor.from_string(MUTED)

    doc.add_paragraph()
    add_table(
        doc,
        ["项目项", "内容"],
        [
            ["当前版本", "Round 46 compendium, 2026-06-20"],
            ["研究定位", "外接式 failure-aware / risk-aware supervisor"],
            ["主要平台", "SurRoL + 早期 3D surgical navigation proxy"],
            ["主要任务", "NeedleReachRL-v0；扩展到 NeedlePickRL-v0、GauzeRetrieveRL-v0"],
            ["核心问题", "How can a surgical robot know when not to recover?"],
            ["证据边界", "仿真原型，不是临床部署或真实组织损伤检测证据"],
        ],
        [1.45, 5.05],
    )
    doc.add_page_break()


def add_static_toc(doc: Document) -> None:
    doc.add_heading("阅读导航", level=1)
    add_numbered(
        doc,
        [
            "执行摘要：一句话说明项目做了什么，哪些证据最强。",
            "相较 2026-06-18 旧版文档的升级：从抽象 3D proxy 走到 SurRoL 视觉可靠性监督。",
            "系统架构：视觉输入、策略、风险头、恢复记忆和风险路由如何连接。",
            "实验全览：Stage 1 到 Stage 6 的完成度、采用模块和撤回模块。",
            "核心实验：多 seed、在线 adapter 数据、recovery memory 负结果、保守路由和跨任务广度。",
            "图表与证据库存：哪些图、表、报告、运行目录支撑当前结论。",
            "博士申请表述：可以怎么讲，不能怎么讲，下一步研究计划。",
        ],
    )
    doc.add_page_break()


def add_exec_summary(doc: Document, figs: dict[str, Path]) -> None:
    doc.add_heading("1. 执行摘要", level=1)
    add_callout(
        doc,
        "一句话定位",
        "本项目不是重新发明手术机器人控制器，而是在已有控制/策略外部增加一个可靠性监督层，用来判断什么时候自动执行、什么时候短程恢复、什么时候交给人工复核、什么时候应视为中止候选。",
        LIGHT_BLUE,
    )
    add_para(
        doc,
        "截至 Round 46，项目已经从最初的抽象 3D surgical tool navigation proxy，升级到 SurRoL 仿真平台上的视觉可靠性监督原型。最强证据集中在 NeedleReach learned visual routing：保守风险路由可以减少自动失败，同时保留已经被允许自动处理的成功轨迹。",
    )
    add_bullets(
        doc,
        [
            "已完成：多 seed 验证、SurRoL 视觉 wrapper、learned risk head、recovery memory、风险分流、recovery budget、temporal stagnation、在线 adapter-space 数据采集、跨任务接口检查。",
            "正结果：risk-aware routing 提高选择性自动成功；budget 10 在 20-seed mixed corruption 下将自动失败从 2 降到 0，并保留 6 条自动成功。",
            "负结果：Round 40 adapter-space memory 会危险放行失败；Round 43 在线 KNN memory 过度复核，不能替代旧 augmented memory。",
            "边界：learned visual routing 不能直接从 NeedleReach 泛化到 NeedlePick/GauzeRetrieve；跨任务正结果目前主要来自 rule/phase-aware monitor 链路。",
        ]
    )
    add_figure(doc, figs["stage"], "图 1. 项目阶段总览。绿色代表可作为当前正向证据，黄色代表完成但应作为限制或负结果解释。")


def add_upgrade_section(doc: Document) -> None:
    doc.add_heading("2. 相较 2026-06-18 旧版文档的升级", level=1)
    add_para(
        doc,
        "旧版 compendium 的主线是抽象 3D navigation proxy：用 constraint-conditioned PPO 与 tangent backup controller 展示“靠近目标时也要受约束地恢复”。新版工作没有丢掉这条线，而是把它升级为更贴近手术机器人仿真的 SurRoL 可靠性监督问题。",
    )
    add_table(
        doc,
        ["维度", "旧版 3D proxy", "新版 SurRoL round46"],
        [
            ["平台", "自建抽象 3D 工具导航环境", "SurRoL surgical robotics simulation"],
            ["主要目标", "避障、目标靠近、tangent shielded recovery", "视觉策略外接可靠性监督与风险分流"],
            ["输入", "低维状态与几何约束", "208D render_proprio_vision，含 proprioception + RGB pooled vision"],
            ["失败处理", "恢复控制器与约束预算", "auto_execute / auto_recovery / human_review / abort_candidate"],
            ["证据形态", "成功率、最终距离、预算耗尽", "多 seed 路由、review rate、auto failure、coverage-risk"],
            ["博士定位", "方法原型", "可讲成 failure-aware surgical autonomy 的研究雏形"],
        ],
        [1.2, 2.55, 2.75],
    )
    add_callout(
        doc,
        "升级的关键",
        "项目从“机器人能不能避开并靠近目标”转向“机器人什么时候不应该继续自动恢复”。这更贴近手术安全，因为错误恢复本身可能带来不可逆后果。",
        PALE_GREEN,
    )


def add_architecture(doc: Document, figs: dict[str, Path]) -> None:
    doc.add_heading("3. 系统架构与风险分流", level=1)
    add_figure(doc, figs["pipeline"], "图 2. 外接式可靠性监督管线。主控制器保持为策略输出，监督层负责风险判断与路由。")
    add_table(
        doc,
        ["路由", "触发含义", "实验解释"],
        [
            ["auto_execute", "低风险，策略继续自动执行", "不是所有轨迹都交给恢复，避免过度干预主流程"],
            ["auto_recovery", "风险升高但仍在短程预算内", "对应可恢复的漂移或短暂视觉扰动"],
            ["human_review", "高风险、OOD、长时间停滞或不确定后果", "对应 ECG 式复查/人工判断思想"],
            ["abort_candidate", "不可逆风险候选或危险区接近", "强调不能一味 recovery，而要允许停止候选"],
        ],
        [1.25, 2.1, 3.15],
    )
    add_para(
        doc,
        "这里的“视觉单元”目前不是完整 CNN/RAM/VLM，而是 SurRoL 渲染图像经过 RGB pooling 得到的伪视觉特征。它已经比直接给目标点更真实，因为策略和风险头需要处理视觉观测扰动，但它仍不是临床级视觉理解。",
    )


def add_stage_table(doc: Document) -> None:
    doc.add_heading("4. 阶段完成度总表", level=1)
    add_table(
        doc,
        ["阶段", "目标", "状态", "关键产物", "结论"],
        [
            ["旧版", "3D proxy constraint RL", "完成", "20260618 compendium, prototype/strict figures", "建立最初不稳定性分析和约束恢复框架"],
            ["Stage 1", "锁定当前 SurRoL 基线", "完成", "current_baseline_summary_round41_zh.md", "采用旧 risk head + 旧 augmented memory，撤回坏 memory"],
            ["Stage 2", "采集真实在线 adapter-space 数据", "完成", "surrol_adapter_online_collection_round42_zh.md", "40 episodes, 1913 steps, 1142 high action-gap"],
            ["Stage 3", "重训在线 recovery memory", "完成但负结果", "surrol_online_adapter_recovery_memory_round43_zh.md", "KNN memory 过度复核，不替代旧 memory"],
            ["Stage 4", "保守风险路由", "完成", "surrol_conservative_routing_round44_zh.md", "budget 10 将自动失败 2 -> 0"],
            ["Stage 5", "跨任务广度", "有限完成", "surrol_cross_task_breadth_round45_zh.md", "wrapper 可接入；learned visual 不能直接跨任务迁移"],
            ["Stage 6", "博士申请定位", "完成", "phd_supervisor_summary_round46_zh.md", "形成可信博士课题雏形"],
        ],
        [0.75, 1.35, 0.8, 1.7, 1.9],
    )


def add_experiments(doc: Document, figs: dict[str, Path]) -> None:
    doc.add_heading("5. 核心实验内容", level=1)

    doc.add_heading("5.1 早期 3D proxy 与多任务 stress", level=2)
    add_para(
        doc,
        "旧版工作证明了项目的初始思想：在抽象 3D 环境中，策略不仅要接近目标，还要在危险区域、预算耗尽和扰动下保持可控。它的价值是形成了“约束、恢复、预算、失败类型”的第一版实验语言。",
    )
    for fig_path, caption in [
        (REPORTS / "figures" / "prototype" / "success_mean.png", "图 3. 旧版 prototype success mean。"),
        (REPORTS / "figures" / "strict" / "final_distance_mean.png", "图 4. 旧版 strict setting final distance mean。"),
    ]:
        add_figure(doc, fig_path, caption, width=5.9)

    doc.add_heading("5.2 Stage 1：SurRoL 当前基线锁定", level=2)
    add_table(
        doc,
        ["模块", "当前采用/候选", "路径或证据", "备注"],
        [
            ["Policy", "采用", "runs/surrol_visual_dagger_round31_seed50710/model_dagger_round2.zip", "NeedleReach learned visual 主策略"],
            ["Risk head", "采用", "visual_action_risk_head.npz", "旧 risk head 是当前主线"],
            ["Recovery memory", "采用", "visual_recovery_memory_augmented.npz", "旧 augmented memory 当前最好"],
            ["Temporal stagnation", "候选 guard", "temporal_stagnation_head.npz", "可减少部分长恢复失败"],
            ["Visual adapter", "候选预处理", "visual_denoising_adapter.npz", "离线降噪有效，闭环仍需谨慎"],
        ],
        [1.2, 1.0, 2.6, 1.7],
    )

    doc.add_heading("5.3 Stage 2：真实在线 adapter-space 数据采集", level=2)
    add_figure(doc, figs["dataset"], "图 5. Round 42 在线 adapter-space 数据规模。")
    add_table(
        doc,
        ["条件", "Episodes", "Steps", "High action-gap", "High-gap rate", "形状/检查"],
        [
            ["mixed p=0.35", "20", "970", "596", "61.4%", "970 x 208, split pass"],
            ["clean", "20", "943", "546", "57.9%", "943 x 208, split pass"],
            ["合计", "40", "1913", "1142", "59.7%", "208D, split pass"],
        ],
        [1.35, 0.8, 0.8, 1.1, 1.0, 1.45],
    )
    add_callout(
        doc,
        "为什么这一步重要",
        "Round 40 的失败说明，不能把旧 DAgger 观测离线映射到 adapter space 后就当作真实闭环数据。Round 42 改为真实在线采集，避免了数据来源错位。",
        PALE_GREEN,
    )

    doc.add_heading("5.4 Stage 3：在线 adapter recovery memory，负结果", level=2)
    add_figure(doc, figs["memory"], "图 6. Recovery memory 闭环对比。旧 memory 仍优于新 KNN adapter memory。")
    add_table(
        doc,
        ["Recovery memory", "成功", "复核", "自动覆盖", "自动成功", "自动失败", "结论"],
        [
            ["旧 augmented memory", "6/20", "12/20", "8/20", "6", "2", "当前保留"],
            ["Round 40 bad new memory", "4/20", "7/20", "13/20", "4", "9", "撤回，危险放行过多"],
            ["Round 43 online KNN", "2/20", "18/20", "2/20", "2", "0", "安全但过度复核，不替代"],
        ],
        [1.55, 0.65, 0.65, 0.8, 0.8, 0.75, 1.3],
    )

    doc.add_heading("5.5 Stage 4：保守风险路由", level=2)
    add_figure(doc, figs["routing"], "图 7. Conservative routing 对比。budget 10 与 stagnation 都能在本组 seed 中把自动失败降到 0。")
    add_table(
        doc,
        ["条件", "成功", "复核", "自动覆盖", "自动成功", "自动失败", "覆盖内成功率", "平均终距"],
        [
            ["baseline old risk + old memory", "6/20", "12/20", "8/20", "6", "2", "75.0%", "0.0882"],
            ["+ recovery budget 10", "6/20", "14/20", "6/20", "6", "0", "100.0%", "0.0823"],
            ["+ learned stagnation", "6/20", "14/20", "6/20", "6", "0", "100.0%", "0.0866"],
            ["+ budget 10 + stagnation", "6/20", "14/20", "6/20", "6", "0", "100.0%", "0.0823"],
            ["new risk + old memory", "4/20", "16/20", "4/20", "4", "0", "100.0%", "0.1180"],
        ],
        [1.55, 0.55, 0.55, 0.7, 0.7, 0.65, 0.85, 0.75],
    )
    add_callout(
        doc,
        "当前推荐主线",
        "old visual risk head + old augmented recovery memory + strict-split visual adapter + selective_memory_guarded + review threshold 0.6 + recovery budget 10。",
        LIGHT_BLUE,
    )

    doc.add_heading("5.6 Stage 5：跨任务广度", level=2)
    add_figure(doc, figs["cross"], "图 8. learned visual 直接迁移失败，但 rule/phase-aware monitor 有跨任务正结果。")
    add_table(
        doc,
        ["任务/链路", "结果", "解释"],
        [
            ["NeedlePick visual wrapper", "obs 208, action 5, smoke pass", "接口不是瓶颈"],
            ["GauzeRetrieve visual wrapper", "obs 208, action 5, smoke pass", "接口不是瓶颈"],
            ["NeedlePick learned transfer", "0/5 success, 5/5 review", "不能说 learned visual 已跨任务成功"],
            ["GauzeRetrieve learned transfer", "0/5 success, 5/5 review", "保守 guard 没有盲目放行"],
            ["NeedlePick phase-aware jaw stuck", "0/10 -> 10/10 recovered", "规则/phase-aware 监督有正结果"],
            ["GauzeRetrieve phase-aware jaw stuck", "0/10 -> 10/10 recovered", "规则/phase-aware 监督有正结果"],
        ],
        [2.0, 1.75, 2.75],
    )


def add_existing_figures(doc: Document) -> None:
    doc.add_heading("6. 已有可视化资产索引", level=1)
    add_para(
        doc,
        "下面这些是项目中已经生成、可继续用于汇报或论文草稿的代表性图。本文档没有把每一张 rollout 轨迹都塞进正文，以免掩盖主线；但关键目录和代表图已经列出。",
    )
    existing = [
        (REPORTS / "figures" / "surrol_risk_triage" / "surrol_risk_triage_routes.png", "图 9. SurRoL risk triage route 分布。"),
        (REPORTS / "figures" / "surrol_reliability_memory" / "embedding_by_route.png", "图 10. Reliability memory embedding 按 route 可视化。"),
        (REPORTS / "figures" / "surrol_phase_aware" / "success_rate_by_failure.png", "图 11. Phase-aware recovery 按失败类型的成功率。"),
        (REPORTS / "figures" / "surrol_severity_sweep" / "needlepick_severity_sweep.png", "图 12. NeedlePick severity sweep。"),
        (REPORTS / "figures" / "surrol_severity_sweep" / "gauzeretrieve_severity_sweep.png", "图 13. GauzeRetrieve severity sweep。"),
    ]
    for p, c in existing:
        add_figure(doc, p, c, width=5.85)
    add_table(
        doc,
        ["目录", "内容"],
        [
            ["reports/figures/prototype", "早期 prototype 成功率、终距、cost、budget exhausted"],
            ["reports/figures/strict", "严格设置下的旧版 3D proxy 指标"],
            ["reports/figures/stress", "多 proxy stress 任务图"],
            ["reports/figures/surrol_risk_triage", "SurRoL route/risk triage 可视化"],
            ["reports/figures/surrol_reliability_memory", "embedding 与 memory route 可视化"],
            ["reports/figures/surrol_phase_aware", "phase-aware recovery 曲线和成功率"],
            ["reports/figures/round46_compendium", "本文档新生成的总览图和对比图"],
        ],
        [2.2, 4.3],
    )


def add_claims(doc: Document) -> None:
    doc.add_heading("7. 结论、限制与可对外表述", level=1)
    add_table(
        doc,
        ["说法", "证据等级", "证据", "限制"],
        [
            ["可靠性监督能提高 NeedleReach 选择性自动执行", "Confirmed", "50-seed guarded routing 优于 pure policy 的 selected success", "主要是 NeedleReach learned visual 设置"],
            ["budget 10 能降低自动失败", "Confirmed", "20-seed mixed corruption 中 auto fail 2 -> 0，auto success 保持 6", "仍需更多任务验证"],
            ["temporal stagnation 能识别部分恢复停滞", "Confirmed but preliminary", "50-seed auto fail 8 -> 4", "训练失败样本仍少"],
            ["visual adapter 离线降噪有效", "Confirmed offline", "strict held-out corruption MSE 降低 99.91%", "不等于闭环恢复有效"],
            ["learned visual 直接跨任务迁移", "Not proven", "NeedlePick/GauzeRetrieve 0/5 success, 5/5 review", "需任务专属 policy/risk/memory"],
            ["临床级组织损伤检测或部署", "Not proven", "当前全是仿真 proxy 风险", "无真实力反馈、组织形变、临床验证"],
            ["Round 40 adapter memory 是正结果", "Withdrawn", "闭环自动失败大幅增加", "只能作为负结果"],
        ],
        [1.85, 1.15, 2.25, 1.25],
    )
    add_callout(
        doc,
        "推荐表述",
        "The current evidence supports a simulation-based external reliability supervisor for selective surgical autonomy. The strongest result is improved routing: reducing unsafe automatic recovery while preserving accepted autonomous successes.",
        PALE_GREEN,
    )
    add_callout(
        doc,
        "避免表述",
        "不要说已经解决 SurRoL 手术机器人自主控制，也不要说 visual adapter 已经解决视觉不确定性，更不要说当前 learned visual supervisor 已跨任务泛化。",
        PALE_RED,
    )


def add_phd_section(doc: Document) -> None:
    doc.add_heading("8. 博士申请 Demo 是否够用", level=1)
    add_para(
        doc,
        "作为博士申请 demo，本项目已经比“我复现了一个环境”更强，因为它有明确的问题定义、正结果、负结果、限制和下一步研究问题。它还不能作为“完整论文级系统”，但足以支撑一个可信的博士课题雏形。",
    )
    add_table(
        doc,
        ["老师会关心的问题", "当前回答"],
        [
            ["你做了什么？", "做了一个外接式可靠性监督层，处理视觉不确定性下的自动执行、恢复、复核和中止候选。"],
            ["为什么不是只做避让？", "现在不是单一避让动作，而是路由框架：风险头、恢复记忆、停滞检测、预算控制、多任务接口。"],
            ["有什么证据？", "多 seed NeedleReach learned visual routing，budget 10 自动失败 2 -> 0；跨任务 rule/phase-aware 有 NeedlePick/GauzeRetrieve 正结果。"],
            ["有什么失败？", "在线 KNN memory 和 adapter-space memory 不能作为正结果；learned visual 不能直接跨任务迁移。"],
            ["博士还能做什么？", "任务专属 visual risk/memory、learned phase-aware recovery head、RAM/VLM embedding、不可逆风险标签、coverage-risk curves。"],
        ],
        [1.95, 4.55],
    )
    add_bullets(
        doc,
        [
            "中文题目建议：面向手术机器人的外接式可靠性监督：视觉不确定性、失败恢复与风险分流。",
            "英文题目建议：Failure-Aware Reliability Supervision for Surgical Robot Autonomy under Visual Uncertainty。",
            "核心研究问题：How can a surgical robot know when not to recover?",
        ]
    )


def add_appendix(doc: Document) -> None:
    doc.add_heading("9. 复现实验与产物清单", level=1)
    add_table(
        doc,
        ["类别", "路径/文件", "用途"],
        [
            ["总阶段文件", "reports/stage_goals_round41.yaml", "记录各阶段完成状态和采用/撤回模块"],
            ["基线总结", "reports/current_baseline_summary_round41_zh.md", "锁定 policy/risk/memory/adapter 主线"],
            ["在线数据", "runs/surrol_visual_adapter_online_mixed20_seed53100", "mixed p=0.35 online adapter data"],
            ["在线数据", "runs/surrol_visual_adapter_online_clean20_seed53200", "clean online adapter data"],
            ["负结果", "reports/surrol_online_adapter_recovery_memory_round43_zh.md", "说明新 KNN memory 不替代旧 memory"],
            ["保守路由", "runs/surrol_visual_round44_conservative_routing_summary.csv", "budget/stagnation/risk head ablation"],
            ["跨任务", "reports/surrol_cross_task_breadth_round45_zh.md", "NeedlePick/GauzeRetrieve 接口和迁移 probe"],
            ["申请叙事", "reports/phd_supervisor_summary_round46_zh.md", "中文给老师看的研究定位"],
            ["英文种子", "reports/phd_proposal_seed_round46_en.md", "英文 proposal seed"],
            ["限制表", "reports/claims_limitations_round46.md", "claims/limitations 分级"],
        ],
        [1.1, 2.6, 2.8],
    )
    add_callout(
        doc,
        "复现提醒",
        "长训练或新 RL 训练仍应先做 smoke run。当前最稳主线是 supervisor/routing/recovery 的仿真实验，而不是声称已训练出强 PPO/RL policy。",
        PALE_GOLD,
    )


def build() -> Path:
    figs = make_figures()
    doc = Document()
    style_doc(doc)
    add_cover(doc)
    add_static_toc(doc)
    add_exec_summary(doc, figs)
    add_upgrade_section(doc)
    add_architecture(doc, figs)
    add_stage_table(doc)
    add_experiments(doc, figs)
    add_existing_figures(doc)
    add_claims(doc)
    add_phd_section(doc)
    add_appendix(doc)
    doc.save(OUT)
    return OUT


if __name__ == "__main__":
    print(build())
