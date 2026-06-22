# -*- coding: utf-8 -*-
"""Build a complete enhanced compendium by preserving the full 2026-06-18
front-experiment DOCX and appending later local evidence.

The intent is conservative: do not rewrite the prior compendium or risk losing
its 46 tables and 48 figures. Instead, treat it as Part I and append an
evidence addendum plus the SurRoL Round 46 synthesis.
"""

from __future__ import annotations

from pathlib import Path

import pandas as pd
from docx import Document

import build_surrol_reliability_compendium_round46 as base


ROOT = Path(__file__).resolve().parents[1]
REPORTS = ROOT / "reports"
RUNS = ROOT / "runs"
SOURCE = REPORTS / "source_full_work_compendium_20260618.docx"
OUT = REPORTS / "constraint_surgical_rl_complete_enhanced_compendium_20260620.docx"


def fmt(x, digits: int = 3) -> str:
    try:
        return f"{float(x):.{digits}f}"
    except Exception:
        return str(x)


def read_agg(path: Path) -> pd.DataFrame:
    return pd.read_csv(path)


def metric(df: pd.DataFrame, variant: str, col: str) -> str:
    row = df.loc[df["variant"] == variant]
    if row.empty:
        return "n/a"
    return fmt(row.iloc[0][col])


def add_source_audit(doc: Document) -> None:
    source_doc = Document(SOURCE)
    base.add_callout(
        doc,
        "增强版原则",
        "Part I 完整保留 2026-06-18 前置实验 DOCX，不重写、不删表、不删图；Part II/III 追加本地后续实验、风险模型与 SurRoL 迁移结果。",
        base.LIGHT_BLUE,
    )
    base.add_table(
        doc,
        ["来源", "本地状态", "在增强版中的处理"],
        [
            [
                "source_full_work_compendium_20260618.docx",
                f"{len(source_doc.paragraphs)} paragraphs, {len(source_doc.tables)} tables, {len(source_doc.inline_shapes)} figures",
                "作为 Part I 原样保留，是前置 3D proxy 实验的最终整合版。",
            ],
            [
                "reports/stress_transfer_suite_report.md",
                "已完成，含 9 个 abstract surgical-proxy preset",
                "补入 Part II，作为多任务前置广度证据。",
            ],
            [
                "runs/failure_suite_report.md",
                "已完成，navigation failure taxonomy",
                "补入 Part II，说明 runtime monitor/recovery 的基础形态。",
            ],
            [
                "runs/manipulation_failure_report.md",
                "已完成，object bias/dropout/slip/contact_loss",
                "补入 Part II，证明不只做 reach/avoid。",
            ],
            [
                "reports/human_review_trigger_report.md",
                "已完成，900 episodes trigger summary",
                "补入 Part II，连接 ECG 式复核/分流思想。",
            ],
            [
                "reports/risk_model_report.md",
                "已完成，proxy + learned risk score",
                "补入 Part II，作为 learned uncertainty/risk head 的前置版本。",
            ],
            [
                "reports/stage_goals_round41.yaml + Round 42-46 reports",
                "已完成 SurRoL 阶段整合",
                "作为 Part III，说明迁移到 SurRoL 后完成与未完成的内容。",
            ],
        ],
        [2.05, 2.15, 2.3],
    )


def add_final_front_version(doc: Document) -> None:
    doc.add_heading("Part II. 本地前置实验增强整合", level=1)
    doc.add_heading("II-1. 前置实验最终版本判断", level=2)
    base.add_para(
        doc,
        "本地文件显示，前置实验不是单个 loose run，而是已经形成一版正式 compendium。增强版采用如下判断：2026-06-18 的 full_work_compendium 是前置 3D proxy 的最终整合版；但同一阶段周边还有 risk/recovery/human-review 文件没有完全进入旧文档，因此作为 Part II 补充。",
    )
    base.add_table(
        doc,
        ["判断项", "结论", "理由"],
        [
            [
                "前置 3D proxy 主报告",
                "以 2026-06-18 full_work_compendium 为准",
                "它已经包含环境、算法、协议、prototype/strict、stress suite、seed-level、48 张图和复现命令。",
            ],
            [
                "5k/20k/50k 训练预算",
                "以本地 aggregate CSV 和旧 compendium 共同为准",
                "pilot_3d_50k 结果与旧文档主结论一致：tangent shield 是主效应。",
            ],
            [
                "failure recovery / human review",
                "旧文档未完整覆盖，按本地 md/csv 追加",
                "reports 与 runs 中存在后续整理文件，尤其是 navigation multi-seed、manipulation failure、risk model。",
            ],
            [
                "SurRoL 迁移",
                "不是前置实验的一部分，作为新阶段",
                "从 6 月 18 日 14:00 后开始形成 SurRoL round 1-46 结果。",
            ],
        ],
        [1.55, 1.65, 3.3],
    )


def add_budget_progression(doc: Document) -> None:
    doc.add_heading("II-2. 3D proxy 训练预算与主效应演化", level=2)
    files = [
        ("5k prototype", RUNS / "pilot_3d_5k_prototype_aggregate_summary.csv"),
        ("20k prototype", RUNS / "pilot_3d_20k_prototype_aggregate_summary.csv"),
        ("50k prototype", RUNS / "pilot_3d_50k_prototype_aggregate_summary.csv"),
    ]
    rows = []
    for label, path in files:
        if not path.exists():
            continue
        df = read_agg(path)
        rows.append(
            [
                label,
                metric(df, "conditioned", "success_mean_mean_over_seeds"),
                metric(df, "conditioned_shielded", "success_mean_mean_over_seeds"),
                metric(df, "conditioned_tangent_shielded", "success_mean_mean_over_seeds"),
                metric(df, "conditioned_tangent_shielded", "budget_exhausted_mean_mean_over_seeds"),
                metric(df, "conditioned_tangent_shielded", "final_distance_mean_mean_over_seeds"),
            ]
        )
    base.add_table(
        doc,
        ["训练预算/版本", "conditioned succ.", "shielded succ.", "tangent succ.", "tangent budget exh.", "tangent final dist."],
        rows,
        [1.3, 1.0, 1.0, 1.0, 1.05, 1.15],
    )
    base.add_callout(
        doc,
        "读法",
        "随着训练预算增加，plain conditioned 仍弱，standard shield 有帮助但不稳定；tangent backup controller 才是贯穿前置实验的主效应。这也解释了为什么后续 SurRoL 阶段转向外接 supervisor/routing，而不是只讲 PPO 本身。",
        base.PALE_GREEN,
    )


def add_stress_summary(doc: Document) -> None:
    doc.add_heading("II-3. 前置多任务 stress transfer suite", level=2)
    stress = [
        ("prototype", RUNS / "stress_prototype_aggregate_summary.csv", "scratch_conditioned_tangent_shielded"),
        ("strict", RUNS / "stress_strict_aggregate_summary.csv", "scratch_conditioned_tangent_shielded"),
        ("needle_reach", RUNS / "stress_needle_reach_aggregate_summary.csv", "scratch_conditioned_tangent_shielded"),
        ("needle_insert", RUNS / "stress_needle_insert_aggregate_summary.csv", "scratch_conditioned_tangent_shielded"),
        ("tight_corridor", RUNS / "stress_tight_corridor_aggregate_summary.csv", "scratch_conditioned_tangent_shielded"),
        ("tissue_retraction_proxy", RUNS / "stress_tissue_retraction_proxy_aggregate_summary.csv", "scratch_conditioned_tangent_shielded"),
        ("gauze_manipulation_proxy", RUNS / "stress_gauze_manipulation_proxy_aggregate_summary.csv", "scratch_conditioned_tangent_shielded"),
        ("peg_transfer_proxy", RUNS / "stress_peg_transfer_proxy_aggregate_summary.csv", "conditioned_tangent_shielded"),
        ("needle_regrasp_proxy", RUNS / "stress_needle_regrasp_proxy_aggregate_summary.csv", "conditioned_tangent_shielded"),
    ]
    rows = []
    for preset, path, variant in stress:
        if path.exists():
            df = read_agg(path)
            rows.append(
                [
                    preset,
                    variant,
                    metric(df, variant, "success_mean_mean_over_seeds"),
                    metric(df, variant, "budget_exhausted_mean_mean_over_seeds"),
                    metric(df, variant, "final_distance_mean_mean_over_seeds"),
                    metric(df, variant, "shield_interventions_mean_mean_over_seeds"),
                ]
            )
    base.add_table(
        doc,
        ["Preset", "Best/current tangent variant", "Success", "Budget exh.", "Final dist.", "Shield int."],
        rows,
        [1.55, 2.05, 0.75, 0.8, 0.85, 0.85],
    )
    base.add_para(
        doc,
        "这部分说明前置实验已经不是单一 reach/avoid：它覆盖 needle reach/insert、tight corridor、tissue retraction、gauze manipulation、peg transfer、needle regrasp 等抽象 surgical proxy。但它仍然是 abstract proxy，不应说成真实 SurRoL 或真实手术机器人验证。",
    )


def add_failure_and_risk_addendum(doc: Document) -> None:
    doc.add_heading("II-4. Failure recovery、人审触发与风险模型", level=2)
    base.add_table(
        doc,
        ["Failure mode", "Controller", "Success", "Detected", "Recovery triggered", "False trigger", "Class correct"],
        [
            ["none", "policy_only", "1.000", "0.000", "0.000", "0.000", "1.000"],
            ["state_target_bias", "policy_only", "0.030", "1.000", "0.000", "0.000", "1.000"],
            ["state_target_bias", "monitor_recovery", "1.000", "1.000", "1.000", "0.000", "1.000"],
            ["state_dropout", "policy_only", "0.030", "1.000", "0.000", "0.000", "1.000"],
            ["state_dropout", "monitor_recovery", "1.000", "1.000", "1.000", "0.000", "1.000"],
            ["execution_slip", "policy_only", "0.100", "1.000", "0.000", "0.000", "1.000"],
            ["execution_slip", "monitor_recovery", "1.000", "1.000", "1.000", "0.000", "1.000"],
        ],
        [1.35, 1.25, 0.75, 0.8, 1.1, 0.9, 0.95],
    )
    base.add_table(
        doc,
        ["Manipulation failure", "Base success", "Monitor success", "Object delivered", "Detection / class"],
        [
            ["none", "1.000", "1.000", "1.000", "no false trigger, class 1.000"],
            ["object_state_bias", "0.010", "1.000", "1.000", "detected 1.000, class 1.000"],
            ["object_dropout", "0.000", "1.000", "1.000", "detected 1.000, class 1.000"],
            ["execution_slip", "0.000", "1.000", "1.000", "detected 1.000, class 1.000"],
            ["contact_loss", "0.000", "1.000", "1.000", "detected 1.000, class 1.000"],
        ],
        [1.55, 1.0, 1.0, 1.05, 1.9],
    )
    base.add_table(
        doc,
        ["Risk/review item", "Episodes", "Main metric", "Meaning"],
        [
            ["Human-review trigger overall", "900", "Precision 1.000, Recall 1.000, False trigger 0.000", "把 monitor 解释为复核触发器，连接 ECG 式分流。"],
            ["Navigation multi-seed", "3 model seeds x 10 episodes/mode", "monitor_recovery abnormal success 1.000", "不是单 checkpoint，增强可靠性证据。"],
            ["Learned risk head proxy overall", "450 held-out", "AUROC 1.000, AUPRC 1.000, ECE 0.001", "前置 learned risk head，仍是合成 proxy failure。"],
            ["Threshold routing", "450 held-out", "threshold 0.2-0.8: precision/recall 1.000", "展示自动覆盖与复核分流思路。"],
        ],
        [1.6, 1.25, 2.15, 1.5],
    )
    base.add_callout(
        doc,
        "限制",
        "这些 failure/review/risk 结果非常整齐，是因为 failure 是合成注入且 monitor 多为规则/仪表化逻辑。它们适合作为博士课题雏形的前置证据，不应包装成真实临床级不确定性模型。",
        base.PALE_GOLD,
    )


def add_surrol_round46_part(doc: Document) -> None:
    doc.add_page_break()
    doc.add_heading("Part III. SurRoL 迁移与 Round 46 完整整合", level=1)
    figs = base.make_figures()
    base.add_exec_summary(doc, figs)
    base.add_upgrade_section(doc)
    base.add_architecture(doc, figs)
    base.add_stage_table(doc)
    base.add_experiments(doc, figs)
    base.add_existing_figures(doc)
    base.add_claims(doc)
    base.add_phd_section(doc)
    base.add_appendix(doc)


def build() -> Path:
    if not SOURCE.exists():
        raise FileNotFoundError(f"Missing source DOCX: {SOURCE}")
    doc = Document(SOURCE)
    doc.add_page_break()
    doc.add_heading("Enhanced Integration Addendum / 增强整合补编", level=1)
    base.add_para(
        doc,
        "本补编生成于 2026-06-20，用于把前置 3D proxy 全量实验、本地后续 failure/risk 文件、以及 SurRoL Round 46 可靠性监督结果合并为一份完整申请材料。",
    )
    add_source_audit(doc)
    add_final_front_version(doc)
    add_budget_progression(doc)
    add_stress_summary(doc)
    add_failure_and_risk_addendum(doc)
    add_surrol_round46_part(doc)
    doc.save(OUT)
    return OUT


if __name__ == "__main__":
    print(build())
