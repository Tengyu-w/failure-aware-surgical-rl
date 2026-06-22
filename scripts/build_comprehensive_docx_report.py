from __future__ import annotations

import csv
from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
RUNS = ROOT / "runs"
REPORTS = ROOT / "reports"
FIGURES = REPORTS / "figures"
OUT = REPORTS / "constraint_surgical_rl_comprehensive_report.docx"

BLUE = RGBColor(46, 116, 181)
DARK_BLUE = RGBColor(31, 77, 120)
INK = RGBColor(18, 32, 48)
GRAY = RGBColor(90, 90, 90)
LIGHT_GRAY = "F4F6F9"
TABLE_HEADER = "E8EEF5"
SOFT_GREEN = "E8F5EC"
SOFT_RED = "FCEAEA"


def read_csv(path: Path) -> list[dict]:
    with path.open(newline="", encoding="utf-8") as f:
        return list(csv.DictReader(f))


def fmt(value: float | str, digits: int = 3) -> str:
    try:
        return f"{float(value):.{digits}f}"
    except (TypeError, ValueError):
        return str(value)


def metric(row: dict, key: str) -> float:
    return float(row[f"{key}_mean_over_seeds"])


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(table, top=80, start=120, bottom=80, end=120) -> None:
    tbl_pr = table._tbl.tblPr
    tbl_cell_mar = tbl_pr.first_child_found_in("w:tblCellMar")
    if tbl_cell_mar is None:
        tbl_cell_mar = OxmlElement("w:tblCellMar")
        tbl_pr.append(tbl_cell_mar)
    for margin, value in {"top": top, "start": start, "bottom": bottom, "end": end}.items():
        node = tbl_cell_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tbl_cell_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_widths(table, widths: list[float]) -> None:
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for row in table.rows:
        for idx, width in enumerate(widths):
            if idx < len(row.cells):
                row.cells[idx].width = Inches(width)


def set_run_font(run, size=None, bold=None, italic=None, color=None, name="Calibri") -> None:
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:ascii"), name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), name)
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    if color is not None:
        run.font.color.rgb = color


def add_paragraph(doc: Document, text: str = "", style: str | None = None, bold=False, italic=False) -> None:
    p = doc.add_paragraph(style=style)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.25
    if text:
        run = p.add_run(text)
        set_run_font(run, bold=bold, italic=italic, color=INK)


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    p = doc.add_paragraph(style=f"Heading {level}")
    run = p.add_run(text)
    if level == 1:
        set_run_font(run, size=16, bold=True, color=BLUE)
        p.paragraph_format.space_before = Pt(18)
        p.paragraph_format.space_after = Pt(10)
    elif level == 2:
        set_run_font(run, size=13, bold=True, color=BLUE)
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after = Pt(6)
    else:
        set_run_font(run, size=12, bold=True, color=DARK_BLUE)
        p.paragraph_format.space_before = Pt(8)
        p.paragraph_format.space_after = Pt(4)


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.line_spacing = 1.208
        run = p.add_run(item)
        set_run_font(run, color=INK)


def add_numbered(doc: Document, items: list[str]) -> None:
    for item in items:
        p = doc.add_paragraph(style="List Number")
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.line_spacing = 1.208
        run = p.add_run(item)
        set_run_font(run, color=INK)


def add_callout(doc: Document, title: str, body: str, fill: str = LIGHT_GRAY) -> None:
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    set_cell_margins(table, top=120, bottom=120, start=160, end=160)
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(title)
    set_run_font(r, size=10.5, bold=True, color=DARK_BLUE)
    p2 = cell.add_paragraph()
    p2.paragraph_format.space_after = Pt(0)
    r2 = p2.add_run(body)
    set_run_font(r2, size=10.5, color=INK)
    doc.add_paragraph()


def add_table(doc: Document, headers: list[str], rows: list[list[str]], widths: list[float] | None = None) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_cell_margins(table)
    hdr = table.rows[0].cells
    for idx, header in enumerate(headers):
        set_cell_shading(hdr[idx], TABLE_HEADER)
        hdr[idx].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p = hdr[idx].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(header)
        set_run_font(run, size=8.5, bold=True, color=INK)
    for row_values in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row_values):
            cells[idx].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            p = cells[idx].paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if idx > 0 else WD_ALIGN_PARAGRAPH.LEFT
            run = p.add_run(value)
            set_run_font(run, size=8.3, color=INK)
    if widths:
        set_table_widths(table, widths)
    doc.add_paragraph()


def add_figure(doc: Document, image_path: Path, caption: str, width: float = 5.9) -> None:
    if not image_path.exists():
        add_callout(doc, "Missing figure", f"Expected figure was not found: {image_path}", fill=SOFT_RED)
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(str(image_path), width=Inches(width))
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.paragraph_format.space_after = Pt(8)
    r = cap.add_run(caption)
    set_run_font(r, size=9, italic=True, color=GRAY)


def configure_styles(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(8)
    normal.paragraph_format.line_spacing = 1.25

    for name in ["Heading 1", "Heading 2", "Heading 3", "List Bullet", "List Number"]:
        style = doc.styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")

    header = section.header.paragraphs[0]
    header.text = "Constraint-Conditioned RL Surgical Tool Navigation | Comprehensive Report"
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    set_run_font(header.runs[0], size=8.5, color=GRAY)

    footer = section.footer.paragraphs[0]
    footer.text = "Generated research artifact | E:\\RL_projects\\constraint_surgical_rl"
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_run_font(footer.runs[0], size=8, color=GRAY)


def cover(doc: Document) -> None:
    doc.add_paragraph()
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("约束条件强化学习用于安全手术工具导航")
    set_run_font(run, size=24, bold=True, color=INK)
    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = p2.add_run("Constraint-Conditioned PPO with Tangent Backup Control")
    set_run_font(r2, size=16, italic=True, color=DARK_BLUE)
    p3 = doc.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r3 = p3.add_run("综合实验报告：数据、细节、规律、图表与复现说明")
    set_run_font(r3, size=14, color=GRAY)
    doc.add_paragraph()
    add_callout(
        doc,
        "Report scope",
        (
            "本报告整合截至当前项目状态的环境设计、算法变体、训练设置、prototype/strict 评估、"
            "seed-level 结果、随机策略 sanity check、轨迹图、规律总结、局限性和下一步计划。"
        ),
        fill=SOFT_GREEN,
    )
    meta = [
        ["Project root", str(ROOT)],
        ["Generated on", date.today().isoformat()],
        ["Main method", "Constraint-conditioned PPO + tangent backup controller"],
        ["Simulator stage", "3D abstract surgical tool navigation proxy"],
        ["Primary report files", "runs/*.csv, reports/figures/*.png, reports/project_brief.md"],
    ]
    add_table(doc, ["Field", "Value"], meta, widths=[1.65, 4.85])
    doc.add_page_break()


def static_toc(doc: Document) -> None:
    add_heading(doc, "目录 / Table of Contents", 1)
    toc_items = [
        "1. Executive Summary / 总结",
        "2. Research Framing / 研究问题",
        "3. Environment And Constraint Design / 环境与约束设计",
        "4. Algorithm Variants / 算法变体",
        "5. Experimental Protocol / 实验协议",
        "6. Main Results / 主要结果",
        "7. Seed-Level Results / 多随机种子结果",
        "8. Figures / 图片与轨迹可视化",
        "9. Experimental Regularities / 实验规律",
        "10. Limitations / 局限",
        "11. Reproduction Commands / 复现命令",
        "12. Data Inventory / 数据文件清单",
        "13. Appendices / 附录",
    ]
    add_numbered(doc, toc_items)
    doc.add_page_break()


def executive_summary(doc: Document, prototype: list[dict], strict: list[dict]) -> None:
    add_heading(doc, "1. Executive Summary / 总结", 1)
    add_paragraph(
        doc,
        (
            "本项目从零搭建了一个抽象 surgical tool navigation 强化学习任务，核心问题是："
            "policy 不仅要到达目标，还要在 task phase 和 safety budget 条件下避免 forbidden tissue-like region。"
            "当前证据表明，单纯 constraint-conditioned PPO 有一定作用，但最强主效应来自 tangent backup controller。"
        ),
    )
    add_callout(
        doc,
        "Main finding / 当前主结论",
        (
            "在 prototype 评估中，scratch/curriculum tangent-shielded policies 均达到 1.000 success、0.000 budget exhaustion、"
            "0.000 cumulative cost；在 strict transfer 中，scratch tangent-shielded 达到 0.860 success 且仍保持 0.000 budget exhaustion。"
            "random policy + tangent shield 的 success 为 0.000，因此 backup controller 并非单独解决 reaching task。"
        ),
        fill=SOFT_GREEN,
    )
    add_heading(doc, "Key Numbers", 2)
    rows = []
    for p_row in sorted(prototype, key=lambda r: r["variant"]):
        s_row = next((r for r in strict if r["variant"] == p_row["variant"]), None)
        rows.append(
            [
                p_row["variant"],
                fmt(metric(p_row, "success_mean")),
                fmt(metric(p_row, "budget_exhausted_mean")),
                fmt(metric(s_row, "success_mean")) if s_row else "-",
                fmt(metric(s_row, "budget_exhausted_mean")) if s_row else "-",
            ]
        )
    add_table(
        doc,
        ["Variant", "Prototype success", "Prototype budget exh.", "Strict success", "Strict budget exh."],
        rows,
        widths=[2.55, 0.95, 0.95, 0.95, 0.95],
    )
    add_heading(doc, "Claim strength", 2)
    add_bullets(
        doc,
        [
            "Confirmed: tangent-shielded methods dominate current prototype safety metrics.",
            "Suggested: tangent backup control can make constraint-conditioned RL more reliable for surgical-tool-style navigation.",
            "Unproven: results have not yet been validated in SurRoL or a high-fidelity surgical simulator.",
        ],
    )
    doc.add_page_break()


def volume_upgrade_context(doc: Document) -> None:
    add_heading(doc, "1A. Volume Upgrade Context / 体量升级参照", 1)
    add_paragraph(
        doc,
        (
            "本轮升级参考了两个已有综合实验文档的结构体量。目标不是简单拉长文字，而是增加实验轴、"
            "压力测试、图像附录、seed-level 表格、复现命令和导师方向映射，使 RL 项目从单一 prototype "
            "报告升级为 compendium-style research artifact。"
        ),
    )
    rows = [
        ["HARP-VLA Comprehensive Experiment Report", "1050", "145", "26", "15", "VLA failure recovery / reliability compendium"],
        ["VT/VF Reliability Compendium v5", "456", "88", "214", "294", "ECG reliability / image-archive compendium"],
        ["RL report before upgrade", "198", "section-limited", "34", "12", "single-task prototype report"],
        ["RL target after upgrade", "stress-suite expanded", "multi-axis", "many aggregate + seed tables", "prototype/strict/stress/rollout figures", "safe surgical tool navigation compendium"],
    ]
    add_table(doc, ["Document", "Paragraphs", "Headings", "Tables", "Figures", "Character"], rows, widths=[2.0, 0.85, 0.8, 0.8, 0.8, 1.25])
    add_heading(doc, "What was missing before this upgrade", 2)
    add_bullets(
        doc,
        [
            "Only one abstract task family was emphasized; now the suite includes needle, corridor, retraction, and gauze proxies.",
            "Only prototype/strict were reported; now there is a seven-preset stress transfer suite.",
            "Only a small image set existed; now every stress preset has success, budget, cost, and distance plots.",
            "The previous document had limited supervisor alignment; now the report explicitly maps experiment axes to target directions.",
        ],
    )
    doc.add_page_break()


def research_framing(doc: Document) -> None:
    add_heading(doc, "2. Research Framing / 研究问题", 1)
    add_paragraph(
        doc,
        (
            "项目原始目标是将普通 peg insertion / PPO benchmark 转化成更贴近手术机器人申请方向的约束操作任务。"
            "我们抽象出 needle reaching/insertion、tissue retraction、gauze manipulation、constrained tool navigation 等任务族，"
            "并优先选择 constrained surgical tool navigation 作为最小可运行研究原型。"
        ),
    )
    add_heading(doc, "Policy interface", 2)
    add_callout(doc, "Policy form", "a_t = pi(s_t, task_phase, safety_budget)")
    add_heading(doc, "Why this is not just a PPO benchmark", 2)
    add_bullets(
        doc,
        [
            "state 中显式包含 task phase 和 remaining safety budget；",
            "奖励之外还跟踪 forbidden-region cost、force proxy、workspace cost 和 budget exhaustion；",
            "方法比较包含 unshielded PPO、standard shield、tangent shield、curriculum 和 random sanity check；",
            "报告 prototype 与 strict transfer 两类评估，而不仅是单一环境分数。",
        ],
    )
    add_heading(doc, "Advisor-facing relevance", 2)
    rows = [
        ["Safe/constrained RL", "Guiliang Liu, Yiding Ji", "constraint budget, backup controller, safe policy learning"],
        ["Surgical robotics", "Fangxun Zhong, Qi Dou, Lu Liu, Hongliang Ren", "surgical tool navigation, forbidden tissue proxy, contact-rich manipulation"],
        ["Robot learning/control", "Changhao Chen, Yuxiang Sun", "policy learning + supervisory correction + transfer check"],
        ["Trustworthy embodied AI", "Minjing Dong, Yang Li", "reliability framing if connected to calibration/transfer/dissertation"],
    ]
    add_table(doc, ["方向", "可能匹配老师", "项目连接点"], rows, widths=[1.55, 2.2, 2.75])
    doc.add_page_break()


def supervisor_direction_mapping(doc: Document) -> None:
    add_heading(doc, "2A. Supervisor-Direction Mapping / 导师方向映射", 1)
    add_paragraph(
        doc,
        (
            "根据已知申请方向与公开研究主题，本项目最适合以 safe surgical robotic manipulation 和 constrained robot learning "
            "来包装。下面的映射用于指导后续实验扩展，而不是替代正式导师主页核查。"
        ),
    )
    rows = [
        ["Fangxun Zhong", "surgical robot autonomy, planning/control", "tool navigation, collision avoidance, task-stage controller", "Very strong"],
        ["Guiliang Liu", "safe/constrained RL, inverse constrained RL", "safety budget, constraint-conditioned PPO, random sanity checks", "Very strong"],
        ["Yiding Ji", "supervisory control, constraints", "tangent backup controller as supervisory safety layer", "Very strong"],
        ["Qi Dou", "surgical embodied autonomy", "surgical-proxy tasks, embodied surgical decision making", "Strong"],
        ["Lu Liu", "AI-driven surgical robotics", "safe surgical manipulation and autonomy", "Strong"],
        ["Hongliang Ren", "medical robotics", "constrained medical tool operation", "Strong"],
        ["Changhao Chen", "robot learning/control", "policy learning plus controller correction", "Medium-strong"],
        ["Yuxiang Sun", "robot learning/control", "manipulation learning and safety-aware policy evaluation", "Medium"],
        ["Minjing Dong", "trustworthy/reliable ML", "robustness and reliability framing only", "Weak-medium"],
        ["Yang Li", "reliability/transfer if framed broadly", "transfer and robustness; dissertation stronger", "Weak-medium"],
    ]
    add_table(doc, ["Supervisor", "Relevant direction", "How this RL project connects", "Fit"], rows, widths=[1.25, 1.65, 2.7, 0.9])
    add_heading(doc, "Implication for experiment design", 2)
    add_bullets(
        doc,
        [
            "For surgical robotics supervisors, expand task proxies and eventually migrate to SurRoL.",
            "For safe RL / supervisory control supervisors, prioritize budget exhaustion, intervention rate, action deviation, and formal ablations.",
            "For trustworthy ML supervisors, connect this project to reliability evidence rather than presenting it as a pure surgical robotics project.",
        ],
    )
    doc.add_page_break()


def environment_design(doc: Document) -> None:
    add_heading(doc, "3. Environment And Constraint Design / 环境与约束设计", 1)
    add_paragraph(
        doc,
        (
            "当前环境是 3D abstract surgical tool navigation proxy。tool tip 从起始体积出发，目标位于目标体积，"
            "中间存在 forbidden tissue-like region。该环境用于快速验证算法想法，后续可迁移到 SurRoL 或 MuJoCo/ManiSkill。"
        ),
    )
    add_heading(doc, "Observation design", 2)
    obs_rows = [
        ["tool_xyz", "3", "工具尖端位置"],
        ["target_xyz", "3", "目标位置"],
        ["forbidden_xyz", "3", "禁区中心"],
        ["distance_to_goal", "1", "目标距离"],
        ["force_proxy", "1", "基于 penetration 的接触/力代理"],
        ["normalized_time", "1", "episode 时间进度"],
        ["task_phase", "1", "approach / align / final 的阶段代理"],
        ["remaining_budget", "1", "剩余安全预算"],
    ]
    add_table(doc, ["Observation component", "Dim", "Meaning"], obs_rows, widths=[2.0, 0.65, 3.85])
    add_heading(doc, "Action and dynamics", 2)
    add_bullets(
        doc,
        [
            "action space: continuous Box(-1, 1), shape=(2,)",
            "tool movement: clipped action scaled by action_scale",
            "workspace: bounded square [-1, 1] x [-1, 1]",
            "termination: success, budget exhaustion, or max step truncation",
        ],
    )
    add_heading(doc, "Preset definitions", 2)
    preset_rows = [
        ["easy", "180", "0.09", "0.10", "2.0-3.0", "curriculum warm-up"],
        ["prototype", "160", "0.07", "0.14", "1.0-2.0", "main training/evaluation"],
        ["strict", "120", "0.055", "0.16", "0.25-0.75", "transfer / harder safety"],
    ]
    add_table(doc, ["Preset", "max steps", "goal radius", "forbidden radius", "budget", "Use"], preset_rows, widths=[0.9, 0.9, 1.0, 1.1, 1.0, 1.6])
    add_heading(doc, "Constraint costs", 2)
    add_bullets(
        doc,
        [
            "forbidden_cost: force_proxy > max_force 时触发；",
            "workspace_cost: tool tip 接近 workspace 边界时触发；",
            "cumulative_cost: forbidden/workspace cost 与 force proxy 的累计安全成本；",
            "budget_exhausted: cumulative_cost 超过 per-episode safety budget。",
        ],
    )
    doc.add_page_break()


def algorithm_variants(doc: Document) -> None:
    add_heading(doc, "4. Algorithm Variants / 算法变体", 1)
    rows = [
        ["conditioned", "PPO + task phase + safety budget", "无 controller", "测试 conditioning 本身是否有用"],
        ["conditioned_shielded", "PPO + standard safety shield", "危险动作改向/停止", "测试简单 backup controller"],
        ["conditioned_tangent_shielded", "PPO + tangent backup controller", "危险动作投影到切向绕行方向", "当前主方法"],
        ["no_phase_budget", "去掉 task phase 和 budget", "无 controller", "普通 baseline"],
        ["curriculum_*", "easy -> prototype", "可叠加 shield", "测试课程学习"],
        ["random_*", "随机动作策略", "可叠加 shield", "sanity check: controller-only 是否解决任务"],
    ]
    add_table(doc, ["Variant", "Policy", "Safety layer", "Purpose"], rows, widths=[1.65, 1.8, 1.55, 1.5])
    add_heading(doc, "Tangent backup controller", 2)
    add_paragraph(
        doc,
        (
            "tangent shield 的设计动机来自可视化检查：standard shield 会让 tool 在 forbidden region 前停住，"
            "安全但不一定能绕行。tangent shield 在 proposed action 进入 forbidden clearance 或越界时，"
            "保留动作的切向分量，并结合 target_tangent 和 radial outward 方向形成安全修正动作。"
        ),
    )
    add_callout(
        doc,
        "Important interpretation",
        (
            "tangent shield 很强，因此需要 random-policy sanity check。当前 random+tangent 的 success=0.000，"
            "说明 controller 单独不解决任务；但它显著降低 unsafe exploration，使 PPO 更容易学习可行路径。"
        ),
        fill=LIGHT_GRAY,
    )
    add_heading(doc, "Curriculum design", 2)
    add_bullets(
        doc,
        [
            "Stage 1: easy preset, wider budget, larger goal radius, smaller forbidden radius。",
            "Stage 2: prototype preset, harder geometry and tighter safety setting。",
            "Same total timesteps are used for scratch vs curriculum comparison where applicable。",
            "当前证据显示 curriculum 有帮助但不是主效应；tangent backup controller 才是主要改进。",
        ],
    )
    doc.add_page_break()


def experimental_protocol(doc: Document) -> None:
    add_heading(doc, "5. Experimental Protocol / 实验协议", 1)
    add_heading(doc, "Training protocol", 2)
    add_bullets(
        doc,
        [
            "Algorithm: Stable-Baselines3 PPO with MlpPolicy.",
            "PPO hyperparameters: n_steps=512, batch_size=128, gamma=0.97, learning_rate default 3e-4.",
            "Main comparison: 50k total timesteps; curriculum uses easy 10k + prototype 40k.",
            "Seeds: 0, 1, 2 for main reported training runs.",
            "Evaluation: deterministic policy, 100 episodes per trained model.",
        ],
    )
    add_heading(doc, "Evaluation metrics", 2)
    metric_rows = [
        ["success", "目标距离小于 goal_radius", "higher is better"],
        ["budget_exhausted", "cumulative_cost > safety_budget", "lower is better"],
        ["cumulative_cost", "episode safety cost", "lower is better"],
        ["final_distance", "episode end 到目标距离", "lower is better"],
        ["final_force_proxy", "final contact/penetration proxy", "lower is better"],
        ["shield_interventions", "backup controller intervention count", "context-dependent"],
    ]
    add_table(doc, ["Metric", "Definition", "Direction"], metric_rows, widths=[1.6, 3.4, 1.5])
    add_heading(doc, "Experiment families", 2)
    rows = [
        ["Smoke tests", "1k-2k timesteps", "check code paths and SB3 compatibility"],
        ["Prototype comparison", "10k x 3 variants x 3 seeds", "early pipeline sanity"],
        ["50k curriculum comparison", "scratch vs easy->prototype", "main prototype evidence"],
        ["Strict transfer", "trained prototype models evaluated on strict", "robustness under tighter constraints"],
        ["Random sanity check", "random actions with/without shield", "controller-only check"],
        ["Rollout visualization", "policy trajectories", "diagnose failure modes and path shape"],
    ]
    add_table(doc, ["Experiment", "Scale", "Purpose"], rows, widths=[1.7, 1.7, 3.1])
    doc.add_page_break()


def main_results(doc: Document, prototype: list[dict], strict: list[dict], random_rows: list[list[str]], heuristic_rows: list[list[str]]) -> None:
    add_heading(doc, "6. Main Results / 主要结果", 1)
    add_heading(doc, "6.1 Prototype aggregate results", 2)
    proto_rows = [
        [
            r["variant"],
            f"{fmt(metric(r, 'success_mean'))} +/- {fmt(r['success_mean_std_over_seeds'])}",
            f"{fmt(metric(r, 'budget_exhausted_mean'))} +/- {fmt(r['budget_exhausted_mean_std_over_seeds'])}",
            f"{fmt(metric(r, 'cumulative_cost_mean'))} +/- {fmt(r['cumulative_cost_mean_std_over_seeds'])}",
            f"{fmt(metric(r, 'final_distance_mean'))} +/- {fmt(r['final_distance_mean_std_over_seeds'])}",
        ]
        for r in sorted(prototype, key=lambda x: x["variant"])
    ]
    add_table(doc, ["Variant", "Success", "Budget exh.", "Cost", "Final dist."], proto_rows, widths=[2.25, 1.1, 1.1, 1.0, 1.0])
    add_heading(doc, "6.2 Strict transfer aggregate results", 2)
    strict_rows = [
        [
            r["variant"],
            f"{fmt(metric(r, 'success_mean'))} +/- {fmt(r['success_mean_std_over_seeds'])}",
            f"{fmt(metric(r, 'budget_exhausted_mean'))} +/- {fmt(r['budget_exhausted_mean_std_over_seeds'])}",
            f"{fmt(metric(r, 'cumulative_cost_mean'))} +/- {fmt(r['cumulative_cost_mean_std_over_seeds'])}",
            f"{fmt(metric(r, 'final_distance_mean'))} +/- {fmt(r['final_distance_mean_std_over_seeds'])}",
        ]
        for r in sorted(strict, key=lambda x: x["variant"])
    ]
    add_table(doc, ["Variant", "Success", "Budget exh.", "Cost", "Final dist."], strict_rows, widths=[2.25, 1.1, 1.1, 1.0, 1.0])
    add_heading(doc, "6.3 Random-policy sanity check", 2)
    add_table(doc, ["Variant", "Success", "Budget exh.", "Cost", "Final dist.", "Shield int."], random_rows, widths=[2.0, 0.8, 0.9, 0.8, 0.9, 0.9])
    add_heading(doc, "6.4 Heuristic environment solvability check", 2)
    add_table(doc, ["Preset", "Success", "Budget exh.", "Cost", "Final dist."], heuristic_rows, widths=[1.3, 1.0, 1.1, 1.0, 1.2])
    add_callout(
        doc,
        "Result pattern",
        (
            "heuristic controller 在 easy/prototype/strict 上都接近或达到高 success，说明环境并非不可解。"
            "random+tangent 不成功，说明 tangent shield 提供安全修正但不是完整任务求解器。"
        ),
        fill=SOFT_GREEN,
    )
    doc.add_page_break()


def stress_suite_section(doc: Document) -> None:
    add_heading(doc, "6A. Expanded Stress Transfer Suite / 扩展压力迁移实验", 1)
    add_paragraph(
        doc,
        (
            "为了把实验体量从单一 prototype 升级到 compendium 级别，本项目新增七个 evaluation presets。"
            "这些 preset 是 surgical-task proxies，用来测试目标精度、预算、禁区半径、运动尺度和接触成本变化下的迁移规律。"
        ),
    )
    add_callout(
        doc,
        "Important caveat",
        (
            "这些 preset 仍然是 3D proxy，不应声称已经等价于真实 needle insertion、tissue retraction 或 gauze manipulation。"
            "它们的作用是形成系统压力测试矩阵，为后续 SurRoL/MuJoCo 迁移提供证据结构。"
        ),
        fill=LIGHT_GRAY,
    )
    stress_files = sorted(RUNS.glob("stress_*_aggregate_summary.csv"))
    for path in stress_files:
        preset = path.name.removeprefix("stress_").removesuffix("_aggregate_summary.csv")
        add_heading(doc, f"Stress preset: {preset}", 2)
        rows = read_csv(path)
        table_rows = []
        for r in sorted(rows, key=lambda x: x["variant"]):
            table_rows.append(
                [
                    r["variant"],
                    fmt(metric(r, "success_mean")),
                    fmt(metric(r, "budget_exhausted_mean")),
                    fmt(metric(r, "cumulative_cost_mean")),
                    fmt(metric(r, "final_distance_mean")),
                    fmt(metric(r, "shield_interventions_mean")),
                ]
            )
        add_table(doc, ["Variant", "Success", "Budget exh.", "Cost", "Distance", "Shield int."], table_rows, widths=[2.3, 0.75, 0.9, 0.75, 0.8, 0.85])
    add_heading(doc, "Cross-preset average pattern", 2)
    report = REPORTS / "stress_transfer_suite_report.md"
    if report.exists():
        add_paragraph(doc, "A markdown stress-suite summary was also generated at reports/stress_transfer_suite_report.md.")
    add_bullets(
        doc,
        [
            "Across seven presets, scratch_conditioned_tangent_shielded achieved the highest mean success among reported variants.",
            "Tangent-shielded variants retained zero observed budget exhaustion and zero cumulative cost across the stress suite.",
            "Standard shield improved safety metrics but did not close the success gap.",
            "Curriculum improved some unshielded settings but did not dominate tangent-shielded scratch training.",
        ],
    )
    doc.add_page_break()


def seed_level_results(doc: Document) -> None:
    add_heading(doc, "7. Seed-Level Results / 多随机种子结果", 1)
    files = [
        ("Prototype all seed-level results", RUNS / "cmp_conditioned_summary.csv"),
        ("Prototype shielded seed-level results", RUNS / "cmp_conditioned_shielded_summary.csv"),
        ("Prototype tangent-shielded seed-level results", RUNS / "cmp_conditioned_tangent_shielded_summary.csv"),
        ("Strict all seed-level results", RUNS / "cmp_strict_all_summary.csv"),
    ]
    for title, path in files:
        add_heading(doc, title, 2)
        rows = read_csv(path)
        table_rows = []
        for r in rows:
            table_rows.append(
                [
                    r["variant"].replace("cmp_", ""),
                    fmt(r["success_mean"]),
                    fmt(r["budget_exhausted_mean"]),
                    fmt(r["cumulative_cost_mean"]),
                    fmt(r["final_distance_mean"]),
                    fmt(r.get("shield_interventions_mean", "0")),
                ]
            )
        add_table(doc, ["Run", "Success", "Budget exh.", "Cost", "Distance", "Shield int."], table_rows, widths=[2.4, 0.75, 0.85, 0.75, 0.8, 0.8])
        doc.add_page_break()


def figures_section(doc: Document) -> None:
    add_heading(doc, "8. Figures / 图片与轨迹可视化", 1)
    add_heading(doc, "8.1 Prototype aggregate plots", 2)
    for name, caption in [
        ("success_mean.png", "Prototype success rate across variants."),
        ("budget_exhausted_mean.png", "Prototype budget exhaustion rate across variants."),
        ("cumulative_cost_mean.png", "Prototype cumulative constraint cost across variants."),
        ("final_distance_mean.png", "Prototype final distance across variants."),
    ]:
        add_figure(doc, FIGURES / "prototype" / name, caption, width=5.7)
    doc.add_page_break()
    add_heading(doc, "8.2 Strict transfer plots", 2)
    for name, caption in [
        ("success_mean.png", "Strict-transfer success rate across variants."),
        ("budget_exhausted_mean.png", "Strict-transfer budget exhaustion rate across variants."),
        ("cumulative_cost_mean.png", "Strict-transfer cumulative constraint cost across variants."),
        ("final_distance_mean.png", "Strict-transfer final distance across variants."),
    ]:
        add_figure(doc, FIGURES / "strict" / name, caption, width=5.7)
    doc.add_page_break()
    add_heading(doc, "8.3 Rollout trajectory figures", 2)
    rollout_paths = sorted((FIGURES / "rollouts").glob("*.png"))
    for idx, path in enumerate(rollout_paths, start=1):
        add_figure(doc, path, f"Tangent-shielded rollout trajectory {idx}.", width=5.55)
        if idx % 2 == 0:
            doc.add_page_break()
    add_heading(doc, "8.4 Stress suite figure appendix", 2)
    stress_root = FIGURES / "stress"
    for preset_dir in sorted(stress_root.glob("*")):
        if not preset_dir.is_dir():
            continue
        add_heading(doc, f"Stress figures: {preset_dir.name}", 3)
        for name, caption in [
            ("success_mean.png", f"{preset_dir.name}: success rate."),
            ("budget_exhausted_mean.png", f"{preset_dir.name}: budget exhaustion."),
            ("cumulative_cost_mean.png", f"{preset_dir.name}: cumulative constraint cost."),
            ("final_distance_mean.png", f"{preset_dir.name}: final distance."),
        ]:
            add_figure(doc, preset_dir / name, caption, width=5.45)
        doc.add_page_break()


def regularities(doc: Document) -> None:
    add_heading(doc, "9. Experimental Regularities / 实验规律", 1)
    patterns = [
        [
            "R1",
            "unshielded PPO 容易耗尽预算",
            "prototype: scratch_conditioned budget exhaustion 0.720; strict: 0.843",
            "单靠 reward penalty 不足以稳定避免 forbidden region。",
        ],
        [
            "R2",
            "standard shield 降低 cost，但 success 提升有限",
            "prototype success 0.313; strict success 0.133",
            "简单 backup controller 更像安全刹车，不一定产生可达路径。",
        ],
        [
            "R3",
            "tangent shield 是主效应",
            "prototype success 1.000; strict scratch tangent success 0.860",
            "切向投影使策略能够绕开禁区，而不仅是停止。",
        ],
        [
            "R4",
            "curriculum 有帮助但不是最强贡献",
            "unshielded prototype success 0.227 -> 0.287",
            "curriculum 可作为辅助训练策略，不宜作为唯一创新点。",
        ],
        [
            "R5",
            "random sanity check 排除了 controller-only 解释",
            "random+tangent success 0.000",
            "backup controller 提供安全修正；到达目标仍依赖 learned policy。",
        ],
        [
            "R6",
            "strict transfer 更能揭示方差",
            "curriculum tangent strict success 0.693 +/- 0.203",
            "下一步需要更多 seeds 和高保真环境验证。",
        ],
    ]
    add_table(doc, ["ID", "Observed regularity", "Evidence", "Interpretation"], patterns, widths=[0.45, 1.75, 2.15, 2.15])
    add_heading(doc, "Failure modes", 2)
    add_bullets(
        doc,
        [
            "早期 standard shield rollout 会在 forbidden region 前停止或 oscillate，安全但难以到达目标。",
            "短训练 tangent shield 可能绕行过度并靠近 workspace 边界；50k 后该问题明显缓解。",
            "strict transfer 下部分 seeds 成功率仍有较大波动，说明 policy robustness 尚不足。",
        ],
    )
    doc.add_page_break()


def limitations_next(doc: Document) -> None:
    add_heading(doc, "10. Limitations And Next Steps / 局限与下一步", 1)
    add_heading(doc, "Limitations", 2)
    add_bullets(
        doc,
        [
            "当前环境是 3D proxy，不是 SurRoL、真实 da Vinci 运动学或 deformable tissue simulation。",
            "主要结果只有 3 个 training seeds，统计证据仍然初步。",
            "tangent shield 很强，需要额外报告 action deviation、intervention timing 和 controller authority。",
            "strict transfer 结果显示方差仍较大，尤其 curriculum tangent-shielded。",
            "目前任务是 constrained navigation，还没有 needle insertion、retraction 或 gauze manipulation 的完整 contact-rich dynamics。",
        ],
    )
    add_heading(doc, "Recommended next experiments", 2)
    add_numbered(
        doc,
        [
            "把同一 policy interface 和 tangent backup controller 接入 SurRoL needle reaching / tool navigation。",
            "增加 action deviation metric：||a_policy - a_executed||，量化 controller 介入强度。",
            "加入 no-controller heuristic controller baseline，进一步隔离 RL policy contribution。",
            "在 strict preset 上做 5-10 seeds，并报告 confidence interval。",
            "加入多阶段任务：approach, align, insert, retract，评估 task phase conditioning 是否真正起效。",
            "生成视频 demo 和 failure case gallery，用于导师邮件和申请 portfolio。",
        ],
    )
    add_heading(doc, "Paper-style wording", 2)
    add_callout(
        doc,
        "Conservative claim",
        (
            "Preliminary results in an abstract surgical navigation task suggest that combining constraint-conditioned PPO "
            "with a tangent backup controller can substantially reduce constraint violations while preserving task success. "
            "These results are promising but remain prototype-level until validated in a higher-fidelity surgical simulator and with additional seeds."
        ),
    )
    doc.add_page_break()


def reproduction(doc: Document) -> None:
    add_heading(doc, "11. Reproduction Commands / 复现命令", 1)
    command_blocks = [
        ("Tests", "& .\\.conda\\python.exe -m pytest -q"),
        ("Prototype comparison", ".\\scripts\\run_prototype_experiment.ps1 -TotalTimesteps 10000 -Episodes 50 -ConfigPreset prototype"),
        ("Curriculum conditioned", ".\\scripts\\run_curriculum_comparison.ps1 -EasyTimesteps 10000 -PrototypeTimesteps 40000 -Episodes 100 -Variant conditioned"),
        ("Curriculum shielded", ".\\scripts\\run_curriculum_comparison.ps1 -EasyTimesteps 10000 -PrototypeTimesteps 40000 -Episodes 100 -Variant conditioned_shielded"),
        ("Curriculum tangent shielded", ".\\scripts\\run_curriculum_comparison.ps1 -EasyTimesteps 10000 -PrototypeTimesteps 40000 -Episodes 100 -Variant conditioned_tangent_shielded"),
        ("Strict transfer", ".\\scripts\\evaluate_trained_comparison.ps1 -EvalPreset strict -Episodes 100"),
        ("Random tangent sanity check", "& .\\.conda\\python.exe scripts\\evaluate_random_policy.py --variant conditioned_tangent_shielded --config-preset prototype --episodes 100"),
        ("Build this report", "& <bundled-python> scripts\\build_comprehensive_docx_report.py"),
    ]
    for title, cmd in command_blocks:
        add_heading(doc, title, 3)
        add_callout(doc, "Command", cmd)
    doc.add_page_break()


def data_inventory(doc: Document) -> None:
    add_heading(doc, "12. Data Inventory / 数据文件清单", 1)
    add_paragraph(
        doc,
        "以下清单列出当前项目中主要结果文件。原始 per-episode CSV 和 monitor logs 保留在 runs/ 目录中；报告只整合关键表格和路径索引。",
    )
    csv_files = sorted(RUNS.rglob("*.csv"))
    rows = []
    for idx, path in enumerate(csv_files, start=1):
        if idx > 160:
            break
        rows.append([str(idx), str(path.relative_to(ROOT)), str(path.stat().st_size)])
    add_table(doc, ["#", "CSV path", "Bytes"], rows, widths=[0.35, 5.3, 0.85])
    if len(csv_files) > 160:
        add_paragraph(doc, f"Note: {len(csv_files) - 160} additional CSV files are present but omitted from this printed inventory for readability.")
    doc.add_page_break()


def appendices(doc: Document) -> None:
    add_heading(doc, "13. Appendices / 附录", 1)
    add_heading(doc, "Appendix A: Important scripts", 2)
    scripts = [
        ["Environment", "src/constraint_surgical_rl/envs/tool_navigation.py", "3D surgical navigation environment and cost logic"],
        ["Wrappers", "src/constraint_surgical_rl/envs/wrappers.py", "standard shield and tangent backup controller"],
        ["Presets", "src/constraint_surgical_rl/envs/presets.py", "easy/prototype/strict config presets"],
        ["Train PPO", "scripts/train_ppo.py", "single-stage PPO training"],
        ["Train curriculum", "scripts/train_curriculum.py", "easy -> prototype training"],
        ["Evaluate policy", "scripts/evaluate_policy.py", "deterministic evaluation metrics"],
        ["Evaluate random", "scripts/evaluate_random_policy.py", "controller-only sanity check"],
        ["Plot rollouts", "scripts/plot_policy_rollouts.py", "trajectory visualization"],
        ["Write reports", "scripts/write_project_brief.py", "research brief generation"],
    ]
    add_table(doc, ["Component", "Path", "Role"], scripts, widths=[1.2, 2.7, 2.6])
    add_heading(doc, "Appendix B: Method-to-claim mapping", 2)
    mapping = [
        ["Constraint-conditioned PPO", "task phase + budget in observation", "policy can react to phase/safety state", "requires ablation in richer environment"],
        ["Standard shield", "lower cost than unshielded", "safety layer helps", "insufficient reaching success"],
        ["Tangent shield", "high success and zero observed budget exhaustion", "safe redirection is strong", "must report intervention authority"],
        ["Curriculum", "small gain in unshielded prototype", "training schedule may help", "not dominant in current results"],
        ["Random sanity check", "success remains 0.000", "controller-only explanation unlikely", "should add heuristic-controller baseline"],
    ]
    add_table(doc, ["Element", "Evidence", "Claim supported", "Caveat"], mapping, widths=[1.2, 1.75, 1.75, 1.8])
    add_heading(doc, "Appendix C: Application positioning", 2)
    add_bullets(
        doc,
        [
            "Best title for portfolio: Constraint-Conditioned PPO with Tangent Backup Control for Safe Surgical Tool Navigation.",
            "Best framing for safe RL supervisors: constrained decision making with supervisory backup control.",
            "Best framing for surgical robotics supervisors: safety-aware surgical tool navigation in forbidden tissue-like geometry.",
            "Best next milestone: SurRoL needle reaching or constrained tool navigation with the same interface.",
        ],
    )


def build() -> None:
    doc = Document()
    configure_styles(doc)
    prototype = read_csv(RUNS / "prototype_all_aggregate_summary.csv")
    strict = read_csv(RUNS / "cmp_strict_all_aggregate_summary.csv")

    random_rows = []
    for variant, path in [
        ("random_conditioned", RUNS / "random_conditioned_prototype.csv"),
        ("random_conditioned_shielded", RUNS / "random_conditioned_shielded_prototype.csv"),
        ("random_conditioned_tangent_shielded", RUNS / "random_conditioned_tangent_shielded_prototype.csv"),
    ]:
        rows = read_csv(path)
        random_rows.append(
            [
                variant,
                fmt(sum(float(r["success"]) for r in rows) / len(rows)),
                fmt(sum(float(r["budget_exhausted"]) for r in rows) / len(rows)),
                fmt(sum(float(r["cumulative_cost"]) for r in rows) / len(rows)),
                fmt(sum(float(r["final_distance"]) for r in rows) / len(rows)),
                fmt(sum(float(r["shield_interventions"]) for r in rows) / len(rows)),
            ]
        )

    heuristic_rows = []
    for preset in ["easy", "prototype", "strict"]:
        path = RUNS / f"heuristic_{preset}_eval.csv"
        rows = read_csv(path)
        heuristic_rows.append(
            [
                preset,
                fmt(sum(float(r["success"]) for r in rows) / len(rows)),
                fmt(sum(float(r["budget_exhausted"]) for r in rows) / len(rows)),
                fmt(sum(float(r["cumulative_cost"]) for r in rows) / len(rows)),
                fmt(sum(float(r["final_distance"]) for r in rows) / len(rows)),
            ]
        )

    cover(doc)
    static_toc(doc)
    executive_summary(doc, prototype, strict)
    volume_upgrade_context(doc)
    research_framing(doc)
    supervisor_direction_mapping(doc)
    environment_design(doc)
    algorithm_variants(doc)
    experimental_protocol(doc)
    main_results(doc, prototype, strict, random_rows, heuristic_rows)
    stress_suite_section(doc)
    seed_level_results(doc)
    figures_section(doc)
    regularities(doc)
    limitations_next(doc)
    reproduction(doc)
    data_inventory(doc)
    appendices(doc)

    REPORTS.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    print(f"docx={OUT}")


if __name__ == "__main__":
    build()
